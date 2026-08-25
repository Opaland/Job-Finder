import io
import re

from fastapi import APIRouter, Depends, HTTPException, Response
from sqlalchemy import or_
from sqlalchemy.orm import Session, load_only

from ..database import get_db
from ..models import OFFER_STATUSES, STATUS_LABELS, Offer, Profile, local_now
from ..schemas import InterviewIn, ManualOffer, OfferDetail, OfferSummary, OfferUpdate
from ..services.claude_ai import ai_cover_letter, ai_email, ai_gap_analysis, ai_interview_prep, cli_available
from ..services.enrich import fetch_full_description
from ..services.journal import log_event
from ..services.scan import find_twin, offer_to_scoring_dict, profile_to_dict, rescore_offer
from ..services.textutils import contains_word, escape_like, fingerprint, normalize

router = APIRouter(prefix="/api/offers", tags=["offres"])

# Colonnes réellement renvoyées par la liste (OfferSummary) : on ne charge pas
# les champs lourds (description, lettres, fiches…) pour les jeter aussitôt.
SUMMARY_COLUMNS = load_only(
    Offer.id, Offer.source, Offer.title, Offer.company, Offer.location,
    Offer.contract_type, Offer.salary_text, Offer.remote, Offer.url,
    Offer.published_at, Offer.collected_at, Offer.still_online, Offer.score,
    Offer.ai_score, Offer.final_score, Offer.status, Offer.favorite,
    Offer.next_action_date,
)


@router.get("", response_model=dict)
def list_offers(
    status: str | None = None,
    source: str | None = None,
    min_score: float | None = None,
    search: str | None = None,
    company: str | None = None,
    favorite: bool | None = None,
    remote: bool | None = None,
    sort: str = "score",
    limit: int = 100,
    offset: int = 0,
    db: Session = Depends(get_db),
):
    query = db.query(Offer)
    if company:
        query = query.filter(Offer.company.ilike(f"%{escape_like(company)}%", escape="\\"))
    if status:
        statuses = [s for s in status.split(",") if s in OFFER_STATUSES]
        if statuses:
            query = query.filter(Offer.status.in_(statuses))
    if source:
        query = query.filter(Offer.source == source)
    if min_score is not None:
        query = query.filter(Offer.final_score >= min_score)
    if favorite is not None:
        query = query.filter(Offer.favorite.is_(favorite))
    if remote is not None:
        query = query.filter(Offer.remote.is_(remote))
    if search:
        like = f"%{escape_like(search)}%"
        query = query.filter(
            or_(
                Offer.title.ilike(like, escape="\\"),
                Offer.company.ilike(like, escape="\\"),
                Offer.description.ilike(like, escape="\\"),
            )
        )

    total = query.count()
    if sort == "date":
        query = query.order_by(Offer.collected_at.desc(), Offer.final_score.desc())
    elif sort == "published":
        query = query.order_by(Offer.published_at.desc().nulls_last(), Offer.final_score.desc())
    else:
        query = query.order_by(Offer.final_score.desc(), Offer.collected_at.desc())
    offers = query.options(SUMMARY_COLUMNS).offset(offset).limit(min(limit, 2000)).all()
    return {
        "total": total,
        "items": [OfferSummary.model_validate(o).model_dump() for o in offers],
    }


@router.post("/manual", response_model=OfferDetail, status_code=201)
def add_manual_offer(payload: ManualOffer, db: Session = Depends(get_db)):
    """Ajoute une offre à la main (annonce collée depuis LinkedIn, Indeed, cooptation…)."""
    text = (payload.raw_text or "").strip()
    if not text and payload.url:
        # Rien de collé : on tente de lire la page de l'offre.
        text = fetch_full_description(payload.url) or ""
    lines = [line.strip() for line in text.splitlines() if line.strip()]

    title = (payload.title or "").strip() or (lines[0][:120] if lines else "")
    if not title:
        raise HTTPException(400, "Donne au moins un titre, ou colle le texte de l'annonce (le titre sera pris sur la première ligne).")
    company = (payload.company or "").strip()
    if not company and len(lines) > 1 and len(lines[1]) <= 60:
        company = lines[1]

    lowered = normalize(text + " " + title)
    contract = ""
    for label, needle in [("CDI", "cdi"), ("CDD", "cdd"), ("Freelance", "freelance"), ("Intérim", "interim"), ("Alternance", "alternance"), ("Stage", "stage")]:
        if contains_word(lowered, needle):
            contract = label
            break

    # Même règle de dédoublonnage que le scan (empreinte, puis titre similaire).
    twin = find_twin(db, title, company)
    if twin:
        raise HTTPException(409, f"Cette offre est déjà suivie : « {twin.title} » ({twin.company or 'entreprise inconnue'}, score {twin.final_score:.0f}).")

    offer = Offer(
        fingerprint=fingerprint(title, company),
        source="manuelle",
        source_id=f"manuelle-{local_now().strftime('%Y%m%d%H%M%S%f')}",
        title=title[:300],
        company=company[:200],
        location=(payload.location or "").strip()[:200],
        description=text,
        url=(payload.url or "").strip(),
        contract_type=contract,
        remote="teletravail" in lowered or "remote" in lowered,
    )
    profile = db.get(Profile, 1)
    rescore_offer(offer, profile_to_dict(profile))
    offer.status_history = [{"status": "nouvelle", "date": local_now().isoformat(), "par": "ajout manuel"}]
    db.add(offer)
    db.commit()
    log_event(db, "ajout", f"Offre ajoutée à la main : « {offer.title} » "
                           f"({offer.company or 'entreprise inconnue'}), score {offer.final_score:.0f}.", offer.id)
    return offer


@router.get("/export.xlsx")
def export_xlsx(db: Session = Depends(get_db)):
    """Exporte toutes les offres suivies dans un classeur Excel (tri par score)."""
    from openpyxl import Workbook
    from openpyxl.styles import Alignment, Font, PatternFill
    from openpyxl.utils import get_column_letter

    headers = [
        ("Score", 8), ("Titre", 42), ("Entreprise", 24), ("Lieu", 20), ("Contrat", 14),
        ("Salaire", 16), ("Télétravail", 11), ("Statut", 12), ("Favori", 8),
        ("Source", 16), ("Publiée le", 12), ("Collectée le", 12), ("Avis IA", 40),
        ("Notes", 40), ("Lien", 45),
    ]

    wb = Workbook()
    ws = wb.active
    ws.title = "Offres"
    header_fill = PatternFill("solid", fgColor="10192B")
    for col, (label, width) in enumerate(headers, start=1):
        cell = ws.cell(row=1, column=col, value=label)
        cell.font = Font(bold=True, color="FFFFFF")
        cell.fill = header_fill
        ws.column_dimensions[get_column_letter(col)].width = width
    ws.freeze_panes = "A2"

    # Seules les colonnes exportées sont lues (pas les descriptions ni les lettres).
    offers = db.query(
        Offer.final_score, Offer.title, Offer.company, Offer.location, Offer.contract_type,
        Offer.salary_text, Offer.remote, Offer.status, Offer.favorite, Offer.source,
        Offer.published_at, Offer.collected_at, Offer.ai_reason, Offer.notes, Offer.url,
    ).order_by(Offer.final_score.desc()).all()
    for row, o in enumerate(offers, start=2):
        values = [
            round(o.final_score), o.title, o.company, o.location, o.contract_type,
            o.salary_text, "Oui" if o.remote else "", STATUS_LABELS.get(o.status, o.status),
            "★" if o.favorite else "",
            o.source, o.published_at.strftime("%d/%m/%Y") if o.published_at else "",
            o.collected_at.strftime("%d/%m/%Y"), o.ai_reason, o.notes, o.url,
        ]
        for col, value in enumerate(values, start=1):
            cell = ws.cell(row=row, column=col, value=value)
            if isinstance(value, str) and value.startswith("="):
                cell.data_type = "s"  # texte collecté, jamais une formule Excel
            cell.alignment = Alignment(vertical="top", wrap_text=col in (2, 13, 14))
    ws.auto_filter.ref = f"A1:{get_column_letter(len(headers))}{max(len(offers) + 1, 2)}"

    buffer = io.BytesIO()
    wb.save(buffer)
    return Response(
        content=buffer.getvalue(),
        media_type="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
        headers={"Content-Disposition": 'attachment; filename="job_finder_offres.xlsx"'},
    )



@router.get("/{offer_id}", response_model=OfferDetail)
def get_offer(offer_id: int, db: Session = Depends(get_db)):
    offer = db.get(Offer, offer_id)
    if not offer:
        raise HTTPException(404, "Offre introuvable")
    return offer


@router.patch("/{offer_id}", response_model=OfferDetail)
def update_offer(offer_id: int, update: OfferUpdate, db: Session = Depends(get_db)):
    offer = db.get(Offer, offer_id)
    if not offer:
        raise HTTPException(404, "Offre introuvable")

    status_change = None
    if update.status is not None:
        if update.status not in OFFER_STATUSES:
            raise HTTPException(400, f"Statut inconnu : {update.status}")
        if update.status != offer.status:
            history = list(offer.status_history or [])
            history.append({"status": update.status, "date": local_now().isoformat(), "par": "utilisateur"})
            offer.status_history = history
            status_change = (offer.status, update.status)
            offer.status = update.status
    if update.notes is not None:
        offer.notes = update.notes
    if update.favorite is not None:
        offer.favorite = update.favorite
    if update.cover_letter is not None:
        offer.cover_letter = update.cover_letter
    if update.interview_prep is not None:
        offer.interview_prep = update.interview_prep
    # Pour la prochaine action, null explicite = effacement.
    if "next_action_date" in update.model_fields_set:
        offer.next_action_date = update.next_action_date
    if "next_action_note" in update.model_fields_set:
        offer.next_action_note = update.next_action_note

    db.commit()
    if status_change:
        log_event(db, "statut", f"« {offer.title} » ({offer.company or 'entreprise inconnue'}) : "
                                f"{status_change[0]} → {status_change[1]}", offer.id)
    return offer


def _generation_ia(db: Session, offer_id: int, generer, *, aide_503: str = "", aide_502: str = "Réessaie dans un instant."):
    """Tronc commun des routes IA : offre existante, CLI présente, résultat non vide.

    `generer(offre_dict, profile)` appelle la fonction ai_* voulue. Renvoie
    (offer, résultat) ; toute évolution de la garde se fait ici une seule fois.
    """
    offer = db.get(Offer, offer_id)
    if not offer:
        raise HTTPException(404, "Offre introuvable")
    if not cli_available():
        raise HTTPException(
            503,
            "CLI Claude Code introuvable sur ce poste. Vérifie que la commande « claude » "
            f"fonctionne dans un terminal{aide_503}.",
        )
    result = generer(offer_to_scoring_dict(offer), db.get(Profile, 1))
    if not result:
        raise HTTPException(502, f"La génération a échoué (voir les logs). {aide_502}")
    return offer, result


@router.post("/{offer_id}/letter", response_model=OfferDetail)
def generate_letter(offer_id: int, db: Session = Depends(get_db)):
    """Génère la lettre de motivation adaptée à l'offre via la session locale Claude Code."""
    offer, letter = _generation_ia(
        db, offer_id,
        lambda o, p: ai_cover_letter(o, p.cv_text, p.letter_template),
        aide_503=", ou édite la lettre manuellement",
        aide_502="Réessaie ou édite la lettre manuellement.",
    )
    offer.cover_letter = letter.strip()
    db.commit()
    log_event(db, "ia", f"Lettre de motivation générée pour « {offer.title} ».", offer.id)
    return offer


@router.post("/{offer_id}/gap-analysis", response_model=OfferDetail)
def generate_gap_analysis(offer_id: int, db: Session = Depends(get_db)):
    """Analyse d'écart CV ↔ offre (compétences couvertes/manquantes, conseils ATS)."""
    offer, analysis = _generation_ia(
        db, offer_id, lambda o, p: ai_gap_analysis(o, p.cv_text)
    )
    offer.gap_analysis = analysis.strip()
    db.commit()
    log_event(db, "ia", f"Analyse d'écart CV ↔ offre générée pour « {offer.title} ».", offer.id)
    return offer


@router.post("/{offer_id}/email")
def generate_email(offer_id: int, kind: str = "candidature", db: Session = Depends(get_db)):
    """Génère un email de candidature ou de relance via la session locale Claude Code."""
    if kind not in ("candidature", "relance"):
        raise HTTPException(400, "Type d'email inconnu : utilise « candidature » ou « relance ».")
    offer, email = _generation_ia(
        db, offer_id, lambda o, p: ai_email(o, p.cv_text, kind)
    )
    log_event(db, "ia", f"Email de {kind} généré pour « {offer.title} ».", offer.id)
    return email


@router.post("/{offer_id}/interview-prep", response_model=OfferDetail)
def generate_interview_prep(offer_id: int, db: Session = Depends(get_db)):
    """Génère une fiche de préparation d'entretien via la session locale Claude Code."""
    offer, prep = _generation_ia(
        db, offer_id, lambda o, p: ai_interview_prep(o, p.cv_text)
    )
    offer.interview_prep = prep.strip()
    db.commit()
    log_event(db, "ia", f"Fiche d'entretien générée pour « {offer.title} ».", offer.id)
    return offer


@router.post("/{offer_id}/enrich", response_model=OfferDetail)
def enrich_offer(offer_id: int, db: Session = Depends(get_db)):
    """Récupère la description complète depuis la page d'origine de l'offre, puis re-score."""
    offer = db.get(Offer, offer_id)
    if not offer:
        raise HTTPException(404, "Offre introuvable")
    text = fetch_full_description(offer.url)
    if not text:
        raise HTTPException(
            502,
            "Impossible de récupérer la description depuis le site d'origine "
            "(page protégée ou indisponible). Ouvre l'offre sur le site via le bouton dédié.",
        )
    if len(text) <= len(offer.description or ""):
        raise HTTPException(409, "La description actuelle est déjà aussi complète que la page d'origine.")

    offer.description = text
    # L'avis IA portait sur l'ancien extrait : on le retire, le prochain scan
    # (ou la prochaine évaluation) le recalculera sur le texte complet.
    offer.ai_score = None
    offer.ai_reason = ""
    profile = db.get(Profile, 1)
    rescore_offer(offer, profile_to_dict(profile))
    db.commit()
    return offer


def _offre(db: Session, offer_id: int) -> Offer:
    offer = db.get(Offer, offer_id)
    if not offer:
        raise HTTPException(404, "Offre introuvable")
    return offer


@router.post("/{offer_id}/interviews", response_model=OfferDetail, status_code=201)
def add_interview(offer_id: int, payload: InterviewIn, db: Session = Depends(get_db)):
    """Planifie (ou consigne) un entretien pour cette offre."""
    offer = _offre(db, offer_id)
    entretiens = list(offer.interviews or [])
    entretiens.append({
        "date": payload.date.isoformat(),
        "format": payload.format.strip()[:60],
        "interlocuteur": payload.interlocuteur.strip()[:120],
        "notes": payload.notes,
        "compte_rendu": "",
        "ressenti": "",
        "suite": "",
    })
    # Toujours trié par date : l'affichage et « le prochain entretien » en dépendent.
    entretiens.sort(key=lambda e: e.get("date") or "")
    offer.interviews = entretiens
    db.commit()
    log_event(db, "entretien", f"Entretien noté pour « {offer.title} » "
                               f"({offer.company or 'entreprise inconnue'}) le "
                               f"{payload.date.strftime('%d/%m/%Y à %H:%M')}.", offer.id)
    return offer


@router.delete("/{offer_id}/interviews/{index}", response_model=OfferDetail)
def delete_interview(offer_id: int, index: int, db: Session = Depends(get_db)):
    offer = _offre(db, offer_id)
    entretiens = list(offer.interviews or [])
    if not 0 <= index < len(entretiens):
        raise HTTPException(404, "Cet entretien n'existe pas (la liste a peut-être changé).")
    entretiens.pop(index)
    offer.interviews = entretiens
    db.commit()
    return offer


@router.get("/{offer_id}/letter.docx")
def letter_docx(offer_id: int, db: Session = Depends(get_db)):
    """Exporte la lettre de motivation de l'offre au format Word."""
    offer = db.get(Offer, offer_id)
    if not offer:
        raise HTTPException(404, "Offre introuvable")
    if not (offer.cover_letter or "").strip():
        raise HTTPException(400, "Aucune lettre pour cette offre : génère-la ou écris-la d'abord.")

    from docx import Document
    from docx.shared import Pt

    profile = db.get(Profile, 1)
    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)

    header = doc.add_paragraph()
    run = header.add_run(profile.full_name or "")
    run.bold = True
    if profile.email:
        header.add_run(f"\n{profile.email}")

    title = f"Candidature — {offer.title}"
    if offer.company:
        title += f" · {offer.company}"
    doc.add_paragraph(title).runs[0].bold = True
    doc.add_paragraph("")

    for paragraph in offer.cover_letter.split("\n\n"):
        if paragraph.strip():
            doc.add_paragraph(paragraph.strip())

    buffer = io.BytesIO()
    doc.save(buffer)

    safe_company = re.sub(r"[^A-Za-z0-9_-]+", "_", offer.company or "offre").strip("_") or "offre"
    return Response(
        content=buffer.getvalue(),
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        headers={"Content-Disposition": f'attachment; filename="lettre_{safe_company}_{offer.id}.docx"'},
    )
